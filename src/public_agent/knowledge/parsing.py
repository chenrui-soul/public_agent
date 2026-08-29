from __future__ import annotations

import hashlib
import io
import unicodedata
import zipfile
from pathlib import PurePosixPath
from typing import Any

from bs4 import BeautifulSoup
from docx import Document
from pydantic import BaseModel, ConfigDict, Field, field_validator
from pypdf import PdfReader

MAX_DOCUMENT_BYTES = 8 * 1024 * 1024
MAX_DOCUMENT_CHARS = 2_000_000

_DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_MEDIA_EXTENSIONS = {
    "text/plain": {".txt"},
    "text/markdown": {".md", ".markdown"},
    "text/html": {".htm", ".html"},
    "application/pdf": {".pdf"},
    _DOCX_MEDIA_TYPE: {".docx"},
}


class DocumentParseError(ValueError):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code


class DocumentSource(BaseModel):
    model_config = ConfigDict(frozen=True)

    filename: str = Field(min_length=1, max_length=255)
    media_type: str = Field(min_length=1, max_length=200)
    content: bytes = Field(min_length=1, max_length=MAX_DOCUMENT_BYTES)

    @field_validator("filename")
    @classmethod
    def validate_filename(cls, value: str) -> str:
        filename = value.strip()
        forbidden = set('<>:"/\\|?*')
        if (
            not filename
            or filename in {".", ".."}
            or filename.endswith((" ", "."))
            or any(ord(char) < 32 or char in forbidden for char in filename)
        ):
            raise ValueError("filename must be a safe leaf name")
        return filename

    @field_validator("media_type")
    @classmethod
    def normalize_media_type(cls, value: str) -> str:
        normalized = value.split(";", 1)[0].strip().lower()
        if not normalized:
            raise ValueError("media_type must not be blank")
        return normalized


class ParsedDocument(BaseModel):
    model_config = ConfigDict(frozen=True)

    filename: str
    media_type: str
    text: str = Field(min_length=1, max_length=MAX_DOCUMENT_CHARS)
    title: str | None = Field(default=None, max_length=500)
    source_hash: str = Field(pattern=r"^[0-9a-f]{64}$")
    parser_profile: str = Field(min_length=1, max_length=100)
    metadata: dict[str, Any] = Field(default_factory=dict)


class DocumentParser:
    def __init__(
        self,
        *,
        max_bytes: int = MAX_DOCUMENT_BYTES,
        max_chars: int = MAX_DOCUMENT_CHARS,
        max_pdf_pages: int = 500,
        max_docx_entries: int = 2_000,
        max_docx_uncompressed_bytes: int = 32 * 1024 * 1024,
        max_docx_compression_ratio: float = 200.0,
    ) -> None:
        if not 1 <= max_bytes <= MAX_DOCUMENT_BYTES:
            raise ValueError(f"max_bytes must be between 1 and {MAX_DOCUMENT_BYTES}")
        if not 1 <= max_chars <= MAX_DOCUMENT_CHARS:
            raise ValueError(f"max_chars must be between 1 and {MAX_DOCUMENT_CHARS}")
        if not 1 <= max_pdf_pages <= 5_000:
            raise ValueError("max_pdf_pages must be between 1 and 5000")
        if not 1 <= max_docx_entries <= 10_000:
            raise ValueError("max_docx_entries must be between 1 and 10000")
        if max_docx_uncompressed_bytes < 1:
            raise ValueError("max_docx_uncompressed_bytes must be positive")
        if max_docx_compression_ratio < 1:
            raise ValueError("max_docx_compression_ratio must be at least 1")
        self._max_bytes = max_bytes
        self._max_chars = max_chars
        self._max_pdf_pages = max_pdf_pages
        self._max_docx_entries = max_docx_entries
        self._max_docx_uncompressed_bytes = max_docx_uncompressed_bytes
        self._max_docx_compression_ratio = max_docx_compression_ratio

    def parse(self, source: DocumentSource) -> ParsedDocument:
        self.validate_source(source)

        if source.media_type in {"text/plain", "text/markdown"}:
            text = _decode_utf8(source.content)
            title = None
            metadata: dict[str, Any] = {}
            parser_profile = "utf8-text-v1"
        elif source.media_type == "text/html":
            text, title, metadata = _parse_html(source.content)
            parser_profile = "html-text-v1"
        elif source.media_type == "application/pdf":
            text, title, metadata = self._parse_pdf(source.content)
            parser_profile = "pypdf-text-v1"
        else:
            text, title, metadata = self._parse_docx(source.content)
            parser_profile = "python-docx-text-v1"

        normalized = _normalize_text(text)
        if len(normalized) > self._max_chars:
            raise DocumentParseError("text_too_large", "parsed text exceeds the character limit")
        if not normalized:
            raise DocumentParseError(
                "no_extractable_text",
                "document does not contain extractable text",
            )
        return ParsedDocument(
            filename=source.filename,
            media_type=source.media_type,
            text=normalized,
            title=_safe_title(title),
            source_hash=hashlib.sha256(source.content).hexdigest(),
            parser_profile=parser_profile,
            metadata=metadata,
        )

    def validate_source(self, source: DocumentSource) -> None:
        """Validate bounded media metadata without extracting document text."""

        if len(source.content) > self._max_bytes:
            raise DocumentParseError("file_too_large", "document exceeds the byte limit")
        allowed_extensions = _MEDIA_EXTENSIONS.get(source.media_type)
        if allowed_extensions is None:
            raise DocumentParseError(
                "unsupported_media_type",
                "document media type is not supported",
            )
        extension = _filename_extension(source.filename)
        if extension not in allowed_extensions:
            raise DocumentParseError(
                "media_type_mismatch",
                "document extension does not match its media type",
            )

    def _parse_pdf(self, content: bytes) -> tuple[str, str | None, dict[str, Any]]:
        if b"%PDF-" not in content[:1024]:
            raise DocumentParseError("invalid_document", "document is not a valid PDF")
        try:
            reader = PdfReader(io.BytesIO(content), strict=False)
            if reader.is_encrypted:
                raise DocumentParseError(
                    "encrypted_document",
                    "encrypted PDF documents are not supported",
                )
            page_count = len(reader.pages)
            if page_count > self._max_pdf_pages:
                raise DocumentParseError(
                    "page_limit_exceeded",
                    "PDF exceeds the configured page limit",
                )
            pages = tuple(page.extract_text() or "" for page in reader.pages)
            metadata = reader.metadata
        except DocumentParseError:
            raise
        except Exception as exc:
            raise DocumentParseError("invalid_document", "PDF parsing failed") from exc
        title = metadata.title if metadata is not None else None
        return "\n\n".join(pages), title, {"page_count": page_count}

    def _parse_docx(self, content: bytes) -> tuple[str, str | None, dict[str, Any]]:
        self._validate_docx_archive(content)
        try:
            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_rows = [
                "\t".join(cell.text for cell in row.cells)
                for table in document.tables
                for row in table.rows
            ]
            title = document.core_properties.title
        except Exception as exc:
            raise DocumentParseError("invalid_document", "DOCX parsing failed") from exc
        return (
            "\n".join((*paragraphs, *table_rows)),
            title,
            {
                "paragraph_count": len(paragraphs),
                "table_count": len(document.tables),
            },
        )

    def _validate_docx_archive(self, content: bytes) -> None:
        if not zipfile.is_zipfile(io.BytesIO(content)):
            raise DocumentParseError("invalid_document", "document is not a valid DOCX")
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                entries = archive.infolist()
                if len(entries) > self._max_docx_entries:
                    raise DocumentParseError(
                        "archive_limit_exceeded",
                        "DOCX contains too many archive entries",
                    )
                names: set[str] = set()
                total_uncompressed = 0
                for entry in entries:
                    path = PurePosixPath(entry.filename.replace("\\", "/"))
                    if path.is_absolute() or ".." in path.parts:
                        raise DocumentParseError(
                            "unsafe_archive_path",
                            "DOCX contains an unsafe archive path",
                        )
                    if entry.flag_bits & 0x1:
                        raise DocumentParseError(
                            "encrypted_document",
                            "encrypted DOCX entries are not supported",
                        )
                    total_uncompressed += entry.file_size
                    if total_uncompressed > self._max_docx_uncompressed_bytes:
                        raise DocumentParseError(
                            "archive_limit_exceeded",
                            "DOCX exceeds the uncompressed byte limit",
                        )
                    ratio = entry.file_size / max(entry.compress_size, 1)
                    if ratio > self._max_docx_compression_ratio:
                        raise DocumentParseError(
                            "archive_limit_exceeded",
                            "DOCX contains an unsafe compression ratio",
                        )
                    names.add(entry.filename)
                if "word/vbaProject.bin" in names:
                    raise DocumentParseError(
                        "unsupported_document_feature",
                        "macro-enabled Word documents are not supported",
                    )
                if not {"[Content_Types].xml", "word/document.xml"}.issubset(names):
                    raise DocumentParseError("invalid_document", "DOCX structure is incomplete")
        except DocumentParseError:
            raise
        except Exception as exc:
            raise DocumentParseError("invalid_document", "DOCX archive validation failed") from exc


def _filename_extension(filename: str) -> str:
    index = filename.rfind(".")
    return filename[index:].lower() if index > 0 else ""


def _decode_utf8(content: bytes) -> str:
    try:
        text = content.decode("utf-8-sig", errors="strict")
    except UnicodeDecodeError as exc:
        raise DocumentParseError(
            "invalid_encoding",
            "text documents must use valid UTF-8",
        ) from exc
    if "\x00" in text:
        raise DocumentParseError("invalid_text", "text documents must not contain NUL bytes")
    return text


def _parse_html(content: bytes) -> tuple[str, str | None, dict[str, Any]]:
    html = _decode_utf8(content)
    try:
        soup = BeautifulSoup(html, "html.parser")
        title = soup.title.get_text(" ", strip=True) if soup.title is not None else None
        for element in soup.select("script, style, noscript, template"):
            element.decompose()
        root = soup.body if soup.body is not None else soup
        text = root.get_text("\n")
    except Exception as exc:
        raise DocumentParseError("invalid_document", "HTML parsing failed") from exc
    return text, title, {"html_element_count": len(soup.find_all())}


def _normalize_text(value: str) -> str:
    normalized = unicodedata.normalize(
        "NFC",
        value.replace("\r\n", "\n").replace("\r", "\n"),
    )
    lines = [" ".join(line.split()) for line in normalized.split("\n")]
    output: list[str] = []
    previous_blank = True
    for line in lines:
        if line:
            output.append(line)
            previous_blank = False
        elif not previous_blank:
            output.append("")
            previous_blank = True
    return "\n".join(output).strip()


def _safe_title(value: object) -> str | None:
    if not isinstance(value, str):
        return None
    normalized = _normalize_text(value)
    return normalized[:500] if normalized else None
