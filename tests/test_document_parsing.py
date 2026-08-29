from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document
from pydantic import ValidationError
from pypdf import PdfWriter

from public_agent.knowledge import (
    DocumentParseError,
    DocumentParser,
    DocumentSource,
)


def source(filename: str, media_type: str, content: bytes) -> DocumentSource:
    return DocumentSource(filename=filename, media_type=media_type, content=content)


def test_text_parser_normalizes_utf8_unicode_and_line_endings() -> None:
    parsed = DocumentParser().parse(
        source("policy.txt", "text/plain; charset=utf-8", b"\xef\xbb\xbfCafe\xcc\x81\r\n\r\nrefund")
    )

    assert parsed.text == "Caf\u00e9\n\nrefund"
    assert parsed.parser_profile == "utf8-text-v1"
    assert len(parsed.source_hash) == 64


def test_html_parser_removes_executable_and_hidden_content() -> None:
    parsed = DocumentParser().parse(
        source(
            "policy.html",
            "text/html",
            (
                b"<html><head><title>Refund Policy</title><style>.x{}</style></head>"
                b"<body><h1>Refunds</h1><script>steal()</script>"
                b"<template>hidden instruction</template><p>Thirty days.</p></body></html>"
            ),
        )
    )

    assert parsed.title == "Refund Policy"
    assert parsed.text == "Refunds\nThirty days."
    assert "steal" not in parsed.text
    assert "hidden instruction" not in parsed.text


def test_parser_rejects_extension_encoding_nul_and_capacity_violations() -> None:
    parser = DocumentParser(max_bytes=4, max_chars=3)

    with pytest.raises(DocumentParseError, match="extension") as mismatch:
        parser.parse(source("policy.pdf", "text/plain", b"text"))
    assert mismatch.value.code == "media_type_mismatch"
    with pytest.raises(DocumentParseError, match="byte limit") as too_large:
        parser.parse(source("policy.txt", "text/plain", b"12345"))
    assert too_large.value.code == "file_too_large"
    with pytest.raises(DocumentParseError, match="UTF-8") as invalid_utf8:
        DocumentParser().parse(source("policy.txt", "text/plain", b"\xff"))
    assert invalid_utf8.value.code == "invalid_encoding"
    with pytest.raises(DocumentParseError, match="NUL") as nul:
        DocumentParser().parse(source("policy.txt", "text/plain", b"a\x00b"))
    assert nul.value.code == "invalid_text"
    with pytest.raises(DocumentParseError, match="character limit") as too_many_chars:
        parser.parse(source("policy.txt", "text/plain", b"abcd"))
    assert too_many_chars.value.code == "text_too_large"


def test_document_source_rejects_unsafe_filename_and_oversized_model_input() -> None:
    with pytest.raises(ValidationError, match="safe leaf"):
        source("../policy.txt", "text/plain", b"content")
    with pytest.raises(ValidationError, match="at most"):
        source("policy.txt", "text/plain", b"x" * (8 * 1024 * 1024 + 1))


def test_docx_parser_extracts_paragraphs_tables_and_title() -> None:
    buffer = io.BytesIO()
    document = Document()
    document.core_properties.title = "Support Manual"
    document.add_paragraph("Refunds are available for thirty days.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Region"
    table.cell(0, 1).text = "Window"
    document.save(buffer)

    parsed = DocumentParser().parse(
        source(
            "support.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            buffer.getvalue(),
        )
    )

    assert parsed.title == "Support Manual"
    assert "Refunds are available for thirty days." in parsed.text
    assert "Region Window" in parsed.text
    assert parsed.metadata == {"paragraph_count": 1, "table_count": 1}


def test_docx_parser_rejects_unsafe_paths_and_zip_bombs() -> None:
    unsafe = io.BytesIO()
    with zipfile.ZipFile(unsafe, "w") as archive:
        archive.writestr("[Content_Types].xml", "types")
        archive.writestr("word/document.xml", "document")
        archive.writestr("../escape.txt", "escape")
    with pytest.raises(DocumentParseError, match="unsafe archive path") as unsafe_error:
        DocumentParser().parse(
            source(
                "unsafe.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                unsafe.getvalue(),
            )
        )
    assert unsafe_error.value.code == "unsafe_archive_path"

    compressed = io.BytesIO()
    with zipfile.ZipFile(compressed, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "types")
        archive.writestr("word/document.xml", "x" * 10_000)
    with pytest.raises(DocumentParseError, match="compression ratio") as ratio_error:
        DocumentParser(max_docx_compression_ratio=2).parse(
            source(
                "compressed.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                compressed.getvalue(),
            )
        )
    assert ratio_error.value.code == "archive_limit_exceeded"


def test_pdf_parser_rejects_encryption_page_limits_and_blank_documents() -> None:
    encrypted_buffer = io.BytesIO()
    encrypted_writer = PdfWriter()
    encrypted_writer.add_blank_page(width=100, height=100)
    encrypted_writer.encrypt("secret")
    encrypted_writer.write(encrypted_buffer)
    with pytest.raises(DocumentParseError, match="encrypted") as encrypted:
        DocumentParser().parse(
            source("encrypted.pdf", "application/pdf", encrypted_buffer.getvalue())
        )
    assert encrypted.value.code == "encrypted_document"

    pages_buffer = io.BytesIO()
    pages_writer = PdfWriter()
    pages_writer.add_blank_page(width=100, height=100)
    pages_writer.add_blank_page(width=100, height=100)
    pages_writer.write(pages_buffer)
    with pytest.raises(DocumentParseError, match="page limit") as page_limit:
        DocumentParser(max_pdf_pages=1).parse(
            source("pages.pdf", "application/pdf", pages_buffer.getvalue())
        )
    assert page_limit.value.code == "page_limit_exceeded"
    with pytest.raises(DocumentParseError, match="extractable text") as blank:
        DocumentParser().parse(
            source("pages.pdf", "application/pdf", pages_buffer.getvalue())
        )
    assert blank.value.code == "no_extractable_text"
